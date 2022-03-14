"""
There are 4 registered models for 'StarDist2D':

Name                  Alias(es)
────                  ─────────
'2D_versatile_fluo'   'Versatile (fluorescent nuclei)'
'2D_versatile_he'     'Versatile (H&E nuclei)'
'2D_paper_dsb2018'    'DSB 2018 (from StarDist 2D paper)'
'2D_demo'             None
"""

from __future__ import print_function, unicode_literals, absolute_import, division
import sys
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
import skimage as skimage
import imageio
import docx
from glob import glob
from tifffile import imread
from csbdeep.utils import Path, normalize
from csbdeep.io import save_tiff_imagej_compatible
from skimage.io import imsave
from skimage import data, restoration, util
import os
from pathlib import Path
from docx import Document
from docx.text import paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from stardist import random_label_cmap, _draw_polygons, export_imagej_rois
from stardist.models import StarDist2D
from skimage.transform import rescale, resize, downscale_local_mean
from stardist.geometry.geom2d import polygons_to_label
from stardist.utils import polyroi_bytearray
from scipy import ndimage, misc
import pandas as pd
import tensorflow as tf

import ssl
ssl._create_default_https_context = ssl._create_unverified_context

MODEL = '2D_versatile_fluo'
PATH_TO_DATA = '/Path/To/dataset' # ENDS _NOT_ WITH A SLASH!!
name = 'Name'  # Name of the experiment

SHOW_IMAGES = True

matplotlib.rcParams["image.interpolation"] = None


np.random.seed(6)
lbl_cmap = random_label_cmap()

print("--> GPU!")
if tf.test.gpu_device_name() == '':
    print('You do not have GPU access.')
else:
    print('You have GPU access')

# Generating results folder and report (report document)

# folders
Results = os.path.join(PATH_TO_DATA, 'Results')
Path(Results).mkdir(parents=True, exist_ok=True)

Results_wi = os.path.join(Results, 'Results_wi')
Path(Results_wi).mkdir(parents=True, exist_ok=True)

Results_ic = os.path.join(Results, 'Results_ic')
Path(Results_ic).mkdir(parents=True, exist_ok=True)

ROI = os.path.join(Results, 'ROI')
Path(ROI).mkdir(parents=True, exist_ok=True)

Measurement_data = os.path.join(Results, 'Measurement_data')
Path(Measurement_data).mkdir(parents=True, exist_ok=True)

# report
report = Document()
styles = report.styles
style = styles.add_style('Calibri', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Calibri'
style.font.size = Pt(10)

report.add_heading('Report of HC data analysis: ' + name, 0)

# Loading raw data
print("--> Loading data!")
filenames = sorted(glob(f'''{PATH_TO_DATA}/*_ch00.tif'''))
X = list(map(imread, filenames))

filenames_YFP = sorted(glob(f'''{PATH_TO_DATA}/*_ch01.tif'''))
Z = list(map(imread, filenames_YFP))

filenames_SR101 = sorted(glob(f'''{PATH_TO_DATA}/*_ch02.tif'''))
Y = list(map(imread, filenames_SR101))

report.add_paragraph('Raw data loaded.')

#report.add_heading('Image rescaling', 1)

# Rescaling for ch00 (nucleus)
print("--> before huge for loop!")
for i in range(len(filenames)):
    image = X[i]
    path_pref = f'''{Results}/{os.path.basename(filenames[i])}'''  # Path name for saving results
    image1 = resize(image, output_shape=(360, 360), preserve_range=True).astype(
        np.float64)  # Rescaling data to 360 x 360 pix
    # save background corrected and rescaled image
    imsave(path_pref + '_nuc_ds.tif', image1, plugin='tifffile')
    plt.imsave(path_pref + '_nuc_ds.png', image1, cmap='magma')
    report.add_picture(path_pref + '_nuc_ds.png',
                       width=Inches(1), height=Inches(1))
print('Hoechst33342 channel images pre-processed sucessufully.')
report.add_paragraph('Hoechst33342 channel images pre-processed sucessufully.')


# Rescaling for ch01 (YFP)

for i in range(len(filenames_YFP)):
    image = Z[i]
    path_pref_YFP = f'''{Results}/{os.path.basename(filenames_YFP[i])}'''  # Path name for saving results
    image1 = resize(image, output_shape=(360, 360), preserve_range=True).astype(
        np.float16)  # Rescaling data to 360 x 360 pix
    # save background corrected and rescaled image
    imsave(path_pref_YFP + '_YFP_ds.tif', image1, plugin='tifffile')
    plt.imsave(path_pref_YFP + '_YFP_ds.png', image1, cmap='magma')
    report.add_picture(path_pref_YFP + '_YFP_ds.png',
                       width=Inches(1), height=Inches(1))
print('YFP channel images pre-processed sucessufully.')
report.add_paragraph('YFP channel images pre-processed sucessufully.')

# Rescaling for ch02 (SR101)

for i in range(len(filenames_SR101)):
    image = Y[i]
    path_pref_SR101 = f'''{Results}/{os.path.basename(filenames_SR101[i])}'''  # Path name for saving results
    image1 = resize(image, output_shape=(360, 360), preserve_range=True).astype(
        np.float16)  # Rescaling data to 360 x 360 pix
    # save background corrected and rescaled image
    imsave(path_pref_SR101 + '_SR101_ds.tif', image1, plugin='tifffile')
    plt.imsave(path_pref_SR101 + '_SR101_ds.png', image1, cmap='magma')
    report.add_picture(path_pref_SR101 + '_SR101_ds.png',
                       width=Inches(1), height=Inches(1))
print('SR101 channel images pre-processed sucessufully.')
report.add_paragraph('SR101 channel images pre-processed sucessufully.')

# Load rescale data

filenames_ds = sorted(glob(f'''{Results}/*nuc_ds.tif'''))
X_ds = list(map(imread, filenames_ds))

filenames_YFP_ds = sorted(glob(f'''{Results}/*_YFP_ds.tif'''))
Z_ds = list(map(imread, filenames_YFP_ds))

filenames_SR101_ds = sorted(glob(f'''{Results}/*SR101_ds.tif'''))
Y_ds = list(map(imread, filenames_SR101_ds))

report.add_paragraph('Rescaled data loaded.')

# Definition of normalization

n_channel = 1 if X[0].ndim == 2 else X[0].shape[-1]
axis_norm = (0, 1)   # normalize channels independently
# axis_norm = (0,1,2) # normalize channels jointly
if n_channel > 1:
    print("Normalizing image channels %s." %
          ('jointly' if axis_norm is None or 2 in axis_norm else 'independently'))
# show all test images
if SHOW_IMAGES:
    fig, ax = plt.subplots(7, 8, figsize=(16, 16))
    for i, (a, x) in enumerate(zip(ax.flat, X)):
        a.imshow(x if x.ndim == 2 else x[..., 0], cmap='gray')
        a.set_title(i)
    [a.axis('off') for a in ax.flat]
    plt.tight_layout()
None


# Segmentation, mask generation and measurement merged


# Define model and and output

report.add_heading(
    'Segmentation, mask generation and measurement: individual cell measurement', 1)
model = StarDist2D.from_pretrained(MODEL)
THRESHOLD = 1000
measurements = []
measurements_merged = []

for i in range(len(filenames_ds)):
    path_pref = f'''{Results}/{os.path.basename(filenames_ds[i])}'''
    path_pref_ic = f'''{Results_ic}/{os.path.basename(filenames_ds[i])}'''
    path_pref_ROI = f'''{ROI}/{os.path.basename(filenames_ds[i])}'''
    img = normalize(X_ds[i], 1, 99.8, axis=axis_norm)
    save_tiff_imagej_compatible(path_pref_ic + '_norm.tiff', img, axes='YX')
    # Segmentation
    labels, details = model.predict_instances(img)
    # export of imageJ compatible ROIs
    export_imagej_rois(path_pref_ROI + '_roi.zip', details['coord'])
    # export center of ROI as .csv file
    np.savetxt(path_pref + '.csv', details['points'], delimiter=',')
    print('model fitted.')
    print('\t', filenames_ds[i])
    report.add_paragraph(str(i) + '\t' + filenames_ds[i])
    report.add_paragraph('model fitted.')
    plt.figure(figsize=(13, 10))
    img_show = img if img.ndim == 2 else img[..., 0]
    coord, points, prob = details['coord'], details['points'], details['prob']
    plt.subplot(121)
    plt.imshow(img_show, cmap='gray')
    plt.axis('off')
    a = plt.axis()
    _draw_polygons(coord, points, prob, show_dist=True)
    plt.axis(a)
    plt.subplot(122)
    plt.imshow(img_show, cmap='gray')
    plt.axis('off')
    plt.imshow(labels, cmap=lbl_cmap, alpha=0.5)
    plt.tight_layout()
    plt.imsave(path_pref_ic + 'seg.png', img_show, cmap='magma')
    plt.imsave(path_pref_ic + 'labels.png', labels, cmap=lbl_cmap)
    report.add_picture(path_pref_ic + 'seg.png',
                       width=Inches(1), height=Inches(1))
    report.add_picture(path_pref_ic + 'labels.png',
                       width=Inches(1), height=Inches(1))
    # Mask generation
    true_mask = np.zeros_like(Z_ds[i])
    labels_normalized = np.zeros_like(labels)
    for n in range(360):
        for m in range(360):
            if (labels[n, m] > 0):
                labels_normalized[n, m] = 1
            val = Z_ds[i][n, m] * labels_normalized[n, m]
            if (val > THRESHOLD):
                true_mask[n, m] = 1
    plt.matshow(true_mask.astype(np.float32))
    plt.imsave(path_pref_ic + '_mask.png', true_mask, cmap='plasma')
    report.add_paragraph('True ROIs generated successfully')
    report.add_picture(path_pref_ic + '_mask.png',
                       width=Inches(1), height=Inches(1))
    # Measurement from each 'cell'
    values = np.zeros_like(Y_ds[i])
    measurements = []
    for n in range(360):
        for m in range(360):
            val = Y_ds[i][n, m] * true_mask[n, m]
            values[n, m] = val
    plt.matshow(values.astype(np.float32))
    matrix = values.copy()
    plt.imsave(path_pref_ic + '_masked_SR101.png', values, cmap='viridis')
    report.add_paragraph('SR101 image masked successfully')
    report.add_picture(path_pref_ic + '_masked_SR101.png',
                       width=Inches(1), height=Inches(1))
    height = 360
    width = 360
    for point in points:
        accumulator = []

        def fill(p):
            x = p[0]
            y = p[1]
            if x < 360 and y < 360:
                if matrix[x][y] == 0 or [x, y] in accumulator:
                    return
                else:
                    accumulator.append([x, y])
                    neighbors = [[x-1, y], [x+1, y], [x-1, y-1], [x+1,
                                                                  y+1], [x-1, y+1], [x+1, y-1], [x, y-1], [x, y+1]]
                    for n in neighbors:
                        if (0 <= n[0] <= width-1) and (0 <= n[1] <= height-1):
                            fill(n)  # do recursive
        ##
        fill(point)
        avg = 0
        for ac in accumulator:
            avg += matrix[ac[0], ac[1]] / len(accumulator)
        if avg > 0:
            measurements.append((avg, point))
            measurements_merged.append((avg, point, i))
            # Shows individual ROIs - comment out to speed up the process
            #accumulator_array = np.array(accumulator)
            #x, y = accumulator_array.T
            # plt.scatter(x,y)
            # plt.show()
    measurements_array = np.array(measurements)
    pd.DataFrame(measurements_array).to_csv(
        Measurement_data + '/Measurements_ic' + str(i) + '.csv')
    report.add_paragraph('Measurement successful.')
measurements_merged_array = np.array(measurements_merged)
pd.DataFrame(measurements_merged_array).to_csv(
    Measurement_data + '/Measurements_merged_ic.csv')

report.add_paragraph(
    'Finished processing. Results are saved in Measurements folder as .csv files.')
report.save(Measurement_data + '/Report_' + name + '.docx')
